import os
import tempfile
import asyncio
import logging
from io import BytesIO
import docx
import PyPDF2
import openpyxl
import pdfplumber
from typing import Optional, Dict, List, Any
from datetime import datetime
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain.docstore.document import Document

# Настройка логирования
logger = logging.getLogger(__name__)

# Импорт репозиториев для работы с PostgreSQL + pgvector
try:
    from backend.database.init_db import get_vector_repository, get_document_repository
    from backend.database.postgresql.models import Document as PGDocument, DocumentVector
    pgvector_available = True
except ImportError as e:
    print(f"Предупреждение: PostgreSQL модули недоступны: {e}")
    print("DocumentProcessor будет работать в режиме fallback (без pgvector)")
    pgvector_available = False
    get_vector_repository = None
    get_document_repository = None
    PGDocument = None
    DocumentVector = None

class DocumentProcessor:
    def __init__(self):
        print("Инициализируем DocumentProcessor...")
        # Инициализация векторного хранилища с pgvector
        self.documents = []  # Кэш для быстрого доступа (опционально)
        self.doc_names = []
        self.embeddings = None
        self.vectorstore = None  # Теперь это флаг, что pgvector используется
        # Репозитории для работы с PostgreSQL
        self.vector_repo = None
        self.document_repo = None
        # Маппинг filename -> document_id для быстрого доступа
        self.filename_to_id: Dict[str, int] = {}
        # Хранилище информации об уверенности для каждого документа
        # {filename: {"confidence": float, "text_length": int, "file_type": str, "words": [{"word": str, "confidence": float}]}}
        self.confidence_data = {}
        # Хранилище путей к изображениям для мультимодальной модели
        # {filename: {"path": file_path, "minio_object": object_name, "minio_bucket": bucket_name}}
        # или {filename: file_path} для обратной совместимости
        self.image_paths = {}
        # Кэш структуры документов для быстрого доступа ко всем чанкам
        # {doc_name: [{"content": str, "chunk": int}, ...]} - отсортировано по chunk
        self._doc_chunks_cache = {}

        logger.info("DocumentProcessor инициализирован")
        self.init_embeddings()
        self.init_pgvector()
        
        # Логируем финальный статус
        status = self.get_pgvector_status()
        if status["available"] and status["initialized"]:
            logger.info(f"PGVECTOR ГОТОВ К РАБОТЕ")
            logger.info(f"   Документов в системе: {status['documents_count']}")
        elif status["available"]:
            logger.warning(f"PGVECTOR ДОСТУПЕН, НО НЕ ИНИЦИАЛИЗИРОВАН")
            if status.get("error"):
                logger.warning(f"   Ошибка: {status['error']}")
        else:
            logger.warning(f"PGVECTOR НЕДОСТУПЕН")
        
    def init_embeddings(self):
        """Инициализация модели для эмбеддингов"""
        print("Инициализируем модель эмбеддингов...")
        try:
            # Путь к локальной модели - сначала проверяем /app/models (монтируется из ./models)
            # Затем проверяем backend/models (для локальной разработки)
            model_name_local = "paraphrase-multilingual-MiniLM-L12-v2"
            
            # Вариант 1: Модель в /app/models (Docker)
            model_path_docker = os.path.join("/app/models", model_name_local)
            # Вариант 2: Модель в backend/models (локальная разработка)
            model_path_local = os.path.join(
                os.path.dirname(os.path.dirname(__file__)),
                "models",
                model_name_local
            )
            
            # Проверяем, существует ли локальная модель
            if os.path.exists(model_path_docker):
                print(f"Используем локальную модель (Docker): {model_path_docker}")
                model_name = model_path_docker
            elif os.path.exists(model_path_local):
                print(f"Используем локальную модель (локально): {model_path_local}")
                model_name = model_path_local
            else:
                # Fallback на Hugging Face Hub
                print("Локальная модель не найдена, загружаем из Hugging Face Hub...")
                model_name = "sentence-transformers/paraphrase-multilingual-MiniLM-L12-v2"
            
            # ВАЖНО: Используем CPU, так как CUDA может не поддерживать новые GPU
            self.embeddings = HuggingFaceEmbeddings(
                model_name=model_name,
                model_kwargs={'device': 'cpu'}  # Принудительно используем CPU
            )
            print("Модель эмбеддингов успешно загружена (CPU)")
        except Exception as e:
            print(f"Ошибка при загрузке модели эмбеддингов: {str(e)}")
            import traceback
            traceback.print_exc()
            self.embeddings = None
    
    def get_pgvector_status(self) -> Dict[str, Any]:
        """
        Получение статуса pgvector
        
        Returns:
            dict: Статус pgvector с детальной информацией
        """
        status = {
            "available": pgvector_available,
            "initialized": False,
            "repositories_ready": False,
            "vectorstore_active": False,
            "documents_count": 0,
            "vectors_count": 0,
            "error": None
        }
        
        if not pgvector_available:
            status["error"] = "Модули PostgreSQL не импортированы"
            return status
        
        try:
            status["initialized"] = self.vector_repo is not None and self.document_repo is not None
            status["repositories_ready"] = status["initialized"]
            status["vectorstore_active"] = self.vectorstore is True
            status["documents_count"] = len(self.doc_names)
            
            # Пытаемся получить количество векторов из БД
            if self.vector_repo:
                try:
                    # Это асинхронный вызов, но для статуса можно пропустить
                    status["vectors_count"] = "N/A (требует async запрос)"
                except:
                    pass
        except Exception as e:
            status["error"] = str(e)
        
        return status
    
    def init_pgvector(self):
        """Инициализация подключения к pgvector"""
        logger.info("=" * 60)
        logger.info("ИНИЦИАЛИЗАЦИЯ PGVECTOR")
        logger.info("=" * 60)
        
        if not pgvector_available:
            logger.warning("PGVECTOR НЕДОСТУПЕН")
            logger.warning("Модули PostgreSQL не импортированы")
            logger.warning("DocumentProcessor будет работать в ограниченном режиме")
            logger.warning("Для полной функциональности установите PostgreSQL с pgvector")
            logger.info("=" * 60)
            return
        
        try:
            logger.info("🔌 Подключение к pgvector...")
            
            # Получаем репозитории
            try:
                self.vector_repo = get_vector_repository()
                logger.info("✅ VectorRepository получен")
            except RuntimeError as e:
                error_msg = str(e)
                if "не инициализирован" in error_msg or "недоступны" in error_msg:
                    logger.error("❌ PostgreSQL не инициализирован")
                    logger.error(f"   {error_msg}")
                    logger.error("   Убедитесь, что:")
                    logger.error("   1. PostgreSQL запущен и доступен")
                    logger.error("   2. Вызван init_postgresql() перед созданием DocumentProcessor")
                    logger.error("   3. Настройки подключения в .env файле корректны")
                    logger.info("=" * 60)
                    self.vector_repo = None
                    self.document_repo = None
                    self.vectorstore = None
                    return
                else:
                    logger.error(f"❌ Ошибка получения VectorRepository: {error_msg}")
                    raise
            
            try:
                self.document_repo = get_document_repository()
                logger.info("✅ DocumentRepository получен")
            except RuntimeError as e:
                error_msg = str(e)
                logger.error(f"❌ Ошибка получения DocumentRepository: {error_msg}")
                self.vector_repo = None
                self.document_repo = None
                self.vectorstore = None
                logger.info("=" * 60)
                return
            
            # Проверяем наличие расширения pgvector
            logger.info("🔍 Проверка наличия расширения pgvector...")
            pgvector_extension_available = asyncio.run(self._check_pgvector_extension())
            
            if not pgvector_extension_available:
                logger.error("❌ PGVECTOR НЕ УСТАНОВЛЕН В POSTGRESQL")
                logger.error("   Расширение 'vector' отсутствует в базе данных")
                logger.error("   Для установки pgvector:")
                logger.error("   1. Установите pgvector в PostgreSQL (см. README/QUICK_START_POSTGRESQL_PGVECTOR.md)")
                logger.error("   2. Или используйте Docker образ с pgvector: pgvector/pgvector:pg17")
                logger.error("   3. После установки выполните: CREATE EXTENSION vector;")
                logger.warning("⚠️  DocumentProcessor будет работать без персистентного хранения")
                logger.info("=" * 60)
                self.vector_repo = None
                self.document_repo = None
                self.vectorstore = None
                return
            
            # Проверяем работоспособность
            logger.info("🔍 Проверка работоспособности pgvector...")
            is_working = asyncio.run(self._test_pgvector_connection())
            
            if is_working:
                self.vectorstore = True  # Флаг, что pgvector используется
                logger.info("✅ PGVECTOR РАБОТАЕТ КОРРЕКТНО")
                logger.info("   - Подключение к PostgreSQL установлено")
                logger.info("   - Расширение pgvector установлено")
                logger.info("   - Репозитории инициализированы")
                logger.info("   - Векторный поиск доступен")
                
                # Загружаем существующие документы из БД
                self._load_documents_from_db()
            else:
                logger.warning("PGVECTOR НЕ РАБОТАЕТ")
                logger.warning("   Проверка работоспособности не прошла")
                self.vector_repo = None
                self.document_repo = None
                self.vectorstore = None
            
            logger.info("=" * 60)
            
        except RuntimeError as e:
            error_msg = str(e)
            logger.error("ОШИБКА ИНИЦИАЛИЗАЦИИ PGVECTOR")
            if "не инициализирован" in error_msg or "недоступны" in error_msg:
                logger.error(f"   PostgreSQL не инициализирован: {error_msg}")
                logger.error("   Убедитесь, что:")
                logger.error("   1. PostgreSQL запущен и доступен")
                logger.error("   2. Вызван init_postgresql() перед созданием DocumentProcessor")
                logger.error("   3. Настройки подключения в .env файле корректны")
            else:
                logger.error(f"   {error_msg}")
            logger.info("=" * 60)
            self.vector_repo = None
            self.document_repo = None
            self.vectorstore = None
        except Exception as e:
            error_msg = str(e)
            logger.error("ОШИБКА ИНИЦИАЛИЗАЦИИ PGVECTOR")
            logger.error(f"   {error_msg}")
            
            # Проверяем, связана ли ошибка с отсутствием расширения
            if "vector" in error_msg.lower() and ("не существует" in error_msg.lower() or "does not exist" in error_msg.lower()):
                logger.error("   Расширение pgvector не установлено в PostgreSQL")
                logger.error("   Установите pgvector согласно инструкции в README/QUICK_START_POSTGRESQL_PGVECTOR.md")
            
            logger.warning("DocumentProcessor будет работать без персистентного хранения")
            import traceback
            logger.debug(f"Traceback: {traceback.format_exc()}")
            logger.info("=" * 60)
            self.vector_repo = None
            self.document_repo = None
            self.vectorstore = None
    
    async def _check_pgvector_extension(self) -> bool:
        """Проверка наличия расширения pgvector в PostgreSQL"""
        try:
            if not self.vector_repo:
                return False
            
            # Используем подключение из репозитория для проверки расширения
            # Используем отдельную транзакцию для избежания конфликтов
            async with self.vector_repo.db_connection.acquire() as conn:
                # Выполняем проверку в отдельной транзакции
                result = await conn.fetchval(
                    "SELECT EXISTS(SELECT 1 FROM pg_extension WHERE extname = 'vector')"
                )
                return bool(result)
        except Exception as e:
            error_msg = str(e)
            # Если ошибка связана с конфликтом операций, предполагаем, что расширение установлено
            # (так как оно создается при инициализации таблиц в repository.py)
            if "another operation is in progress" in error_msg.lower() or "операция уже выполняется" in error_msg.lower():
                logger.info("Проверка расширения пропущена из-за активной операции, предполагаем что расширение установлено")
                logger.info("(Расширение создается автоматически при инициализации таблиц)")
                return True
            logger.warning(f"Не удалось проверить расширение pgvector: {error_msg}")
            # Если проверка не удалась, но мы знаем, что расширение должно быть установлено
            # (так как оно создается в repository.py), возвращаем True
            logger.info("Предполагаем, что расширение установлено (создается при инициализации таблиц)")
            return True
    
    async def _test_pgvector_connection(self) -> bool:
        """Тестирование подключения к pgvector"""
        try:
            # Проверяем, что репозитории доступны
            if not self.vector_repo or not self.document_repo:
                logger.warning("Репозитории не инициализированы")
                return False
            
            # Пробуем выполнить простой запрос к БД
            # Проверяем количество документов в БД
            try:
                documents = await self.document_repo.get_all_documents(limit=1)
                logger.info(f" В базе данных найдено документов: {len(documents)} (проверка ограничена 1)")
            except Exception as e:
                logger.warning(f"Не удалось проверить документы в БД: {str(e)}")
                return False
            
            # Проверяем, что таблицы векторов существуют
            # Это делается через попытку выполнить простой запрос
            try:
                # Пробуем выполнить поиск с пустым вектором (просто для проверки таблицы)
                test_embedding = [0.0] * 384  # Размерность по умолчанию
                results = await self.vector_repo.similarity_search(test_embedding, limit=1)
                logger.info(f"Таблица векторов доступна, найдено векторов: {len(results)} (проверка ограничена 1)")
            except Exception as e:
                error_msg = str(e)
                # Проверяем, связана ли ошибка с отсутствием типа vector
                if "vector" in error_msg.lower() and ("не существует" in error_msg.lower() or "does not exist" in error_msg.lower()):
                    logger.error(f"Таблица векторов не может быть использована: {error_msg}")
                    logger.error("   Расширение pgvector не установлено или таблица не создана")
                    return False
                else:
                    logger.warning(f"Не удалось проверить таблицу векторов: {error_msg}")
                    # Это не критично, если таблица еще не создана
                    logger.info("   (Таблица векторов может быть пустой - это нормально)")
            
            return True
            
        except Exception as e:
            error_msg = str(e)
            logger.error(f"Ошибка при проверке работоспособности: {error_msg}")
            
            # Проверяем, связана ли ошибка с отсутствием расширения
            if "vector" in error_msg.lower() and ("не существует" in error_msg.lower() or "does not exist" in error_msg.lower()):
                logger.error("   Расширение pgvector не установлено в PostgreSQL")
            
            import traceback
            logger.debug(f"   Traceback: {traceback.format_exc()}")
            return False
    
    def _load_documents_from_db(self):
        """Загрузка документов из базы данных при инициализации"""
        if not self.document_repo:
            logger.warning("DocumentRepository недоступен, пропускаем загрузку документов")
            return
        
        try:
            logger.info("📥 Загрузка документов из базы данных...")
            # Получаем все документы из БД (синхронный вызов async метода)
            pg_documents = asyncio.run(self.document_repo.get_all_documents(limit=1000))
            
            loaded_count = 0
            for pg_doc in pg_documents:
                filename = pg_doc.filename
                if filename not in self.doc_names:
                    self.doc_names.append(filename)
                    loaded_count += 1
                self.filename_to_id[filename] = pg_doc.id
                
                # Загружаем метаданные если есть
                if pg_doc.metadata:
                    if "confidence_data" in pg_doc.metadata:
                        self.confidence_data[filename] = pg_doc.metadata["confidence_data"]
            
            if loaded_count > 0:
                logger.info(f"Загружено {loaded_count} новых документов из базы данных")
                logger.info(f"Всего документов в системе: {len(self.doc_names)}")
            else:
                logger.info("Документы в базе данных уже загружены или база пуста")
        except Exception as e:
            logger.error(f"Ошибка при загрузке документов из БД: {str(e)}")
            import traceback
            logger.debug(f"Traceback: {traceback.format_exc()}")
    
    async def process_document(self, file_data: bytes, filename: str, file_extension: str, minio_object_name=None, minio_bucket=None):
        """
        Обработка документа в зависимости от его типа
        
        Args:
            file_data: Данные файла в виде bytes
            filename: Имя файла (для идентификации)
            file_extension: Расширение файла (например, '.pdf', '.docx')
            minio_object_name: Имя объекта в MinIO (если файл хранится в MinIO)
            minio_bucket: Имя bucket в MinIO (если файл хранится в MinIO)
        """
        file_extension = file_extension.lower()
        document_text = ""
        confidence_info = None
        
        try:
            print(f"Обрабатываем документ: {filename} (тип: {file_extension}, размер: {len(file_data)} байт)")
            
            if file_extension == '.docx':
                result = self.extract_text_from_docx_bytes(file_data)
                if isinstance(result, dict):
                    document_text = result.get("text", "")
                    confidence_info = result.get("confidence_info", {"confidence": 100.0, "text_length": len(document_text), "file_type": "docx", "words": []})
                else:
                    document_text = result
                    confidence_info = self._create_confidence_info_for_text(document_text, 100.0, "docx")
            elif file_extension == '.pdf':
                result = self.extract_text_from_pdf_bytes(file_data)
                if isinstance(result, dict):
                    document_text = result.get("text", "")
                    confidence_info = result.get("confidence_info", {"confidence": 100.0, "text_length": len(document_text), "file_type": "pdf", "words": []})
                else:
                    document_text = result
                    confidence_info = self._create_confidence_info_for_text(document_text, 100.0, "pdf")
            elif file_extension in ['.xlsx', '.xls']:
                result = self.extract_text_from_excel_bytes(file_data)
                if isinstance(result, dict):
                    document_text = result.get("text", "")
                    confidence_info = result.get("confidence_info", {"confidence": 100.0, "text_length": len(document_text), "file_type": "excel", "words": []})
                else:
                    document_text = result
                    confidence_info = self._create_confidence_info_for_text(document_text, 100.0, "excel")
            elif file_extension == '.txt':
                result = self.extract_text_from_txt_bytes(file_data)
                if isinstance(result, dict):
                    document_text = result.get("text", "")
                    confidence_info = result.get("confidence_info", {"confidence": 100.0, "text_length": len(document_text), "file_type": "txt", "words": []})
                else:
                    document_text = result
                    confidence_info = self._create_confidence_info_for_text(document_text, 100.0, "txt")
            elif file_extension in ['.jpg', '.jpeg', '.png', '.webp']:
                result = self.extract_text_from_image_bytes(file_data)
                if isinstance(result, dict):
                    document_text = result.get("text", "")
                    confidence_info = result.get("confidence_info", {"confidence": 0.0, "text_length": len(document_text), "file_type": "image", "words": []})
                    # Если OCR не удался (пустой текст и есть ошибка), логируем, но продолжаем обработку
                    if not document_text and confidence_info.get("error"):
                        error_msg = confidence_info.get("error", "Неизвестная ошибка")
                        print(f"ВНИМАНИЕ: OCR не удался для {filename}: {error_msg}")
                        print(f"Изображение будет сохранено, но текст не будет извлечен")
                        # Продолжаем обработку - документ будет добавлен с пустым текстом
                        # Изображение все равно будет доступно для мультимодальной модели
                else:
                    document_text = result
                    confidence_info = self._create_confidence_info_for_text(document_text, 50.0, "image")
            else:
                return False, f"Неподдерживаемый формат файла: {file_extension}"
            
            print(f"Извлечено текста: {len(document_text)} символов")
            
            # Сохраняем информацию об уверенности
            if confidence_info:
                self.confidence_data[filename] = confidence_info
                print(f"Сохранена информация об уверенности для {filename}: {confidence_info['confidence']:.2f}%")
            
            # Сохраняем информацию об изображении в MinIO, если это изображение
            if file_extension in ['.jpg', '.jpeg', '.png', '.webp']:
                # Сохраняем информацию о MinIO объекте
                if minio_object_name and minio_bucket:
                    self.image_paths[filename] = {
                        "minio_object": minio_object_name,
                        "minio_bucket": minio_bucket,
                        "file_data": file_data  # Сохраняем данные в памяти для быстрого доступа
                    }
                    print(f"Сохранена информация об изображении в MinIO для {filename}: {minio_bucket}/{minio_object_name}")
                else:
                    # Fallback: сохраняем данные в памяти
                    self.image_paths[filename] = {
                        "file_data": file_data
                    }
                    print(f"Сохранены данные изображения в памяти для {filename}")
            
            # Добавляем документ в коллекцию
            # Если текст пустой (например, OCR не сработал), все равно добавляем документ
            # чтобы изображение было доступно для мультимодальной модели
            if document_text or file_extension in ['.jpg', '.jpeg', '.png', '.webp']:
                # Для изображений добавляем даже с пустым текстом
                # Используем async версию, так как process_document теперь async
                await self.add_document_to_collection_async(document_text, filename)
                print(f"Документ добавлен в коллекцию. Всего документов: {len(self.doc_names)}")
            else:
                print(f"Пропускаем добавление документа {filename} - текст пустой и это не изображение")
            
            return True, f"Документ {filename} успешно обработан"
            
        except Exception as e:
            print(f"Ошибка при обработке документа: {str(e)}")
            return False, f"Ошибка при обработке документа: {str(e)}"
    
    def _create_confidence_info_for_text(self, text, confidence_per_word, file_type):
        """Создание структуры информации об уверенности для текста"""
        import re
        # Улучшенное разбиение на слова: разделяем слова и знаки препинания
        # Находим слова (буквы, цифры, дефисы внутри слов) и знаки препинания отдельно
        # Паттерн: \w+ для слов (включая буквы, цифры, подчеркивания), или [^\w\s] для знаков препинания
        # Но лучше использовать более простой подход: разбиваем по пробелам и сохраняем структуру
        
        # Разбиваем текст на токены, сохраняя структуру
        # Используем регулярное выражение, которое находит слова и знаки препинания отдельно
        tokens = re.findall(r'\w+|[^\w\s]+', text)
        
        # Фильтруем пустые токены и формируем список слов
        words_with_confidence = []
        for token in tokens:
            token = token.strip()
            if token:  # Пропускаем пустые токены
                words_with_confidence.append({"word": token, "confidence": float(confidence_per_word)})
        
        avg_confidence = confidence_per_word
        
        return {
            "confidence": avg_confidence,
            "text_length": len(text),
            "file_type": file_type,
            "words": words_with_confidence
        }
    
    def extract_text_from_docx_bytes(self, file_data: bytes):
        """Извлечение текста из DOCX файла из bytes"""
        print(f"Извлекаем текст из DOCX файла (размер: {len(file_data)} байт)")
        doc = docx.Document(BytesIO(file_data))
        full_text = []
        
        for para in doc.paragraphs:
            full_text.append(para.text)
        
        # Извлекаем текст из таблиц
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    full_text.append(cell.text)
        
        result = "\n".join(full_text)
        print(f"Извлечено {len(result)} символов из DOCX")
        return result
    
    def extract_text_from_pdf_bytes(self, file_data: bytes):
        """Извлечение текста из PDF файла из bytes с информацией об уверенности"""
        print(f"Извлекаем текст из PDF файла (размер: {len(file_data)} байт)")
        text = ""
        confidence_scores = []
        total_chars = 0
        
        # Используем PDFPlumber для более точного извлечения текста
        try:
            with pdfplumber.open(BytesIO(file_data)) as pdf:
                for page_num, page in enumerate(pdf.pages):
                    page_text = page.extract_text() or ""
                    text += page_text
                    total_chars += len(page_text)
                    
                    # Для PDF оцениваем уверенность на основе возможности извлечения текста
                    # Если текст извлекается успешно, считаем уверенность высокой
                    if page_text.strip():
                        # PDFPlumber обычно имеет высокую уверенность, если текст извлекается
                        confidence_scores.append(95.0)  # Высокая уверенность для успешного извлечения
                    else:
                        confidence_scores.append(50.0)  # Средняя уверенность, если текст не найден
                        
            print(f"PDFPlumber успешно извлек {len(text)} символов")
        except Exception as e:
            print(f"Ошибка при извлечении текста с помощью pdfplumber: {str(e)}")
            
            # Резервный метод с PyPDF2
            try:
                reader = PyPDF2.PdfReader(BytesIO(file_data))
                for page in reader.pages:
                    page_text = page.extract_text() or ""
                    text += page_text
                    total_chars += len(page_text)
                    if page_text.strip():
                        confidence_scores.append(85.0)  # Немного ниже уверенность для PyPDF2
                    else:
                        confidence_scores.append(40.0)
                print(f"PyPDF2 успешно извлек {len(text)} символов")
            except Exception as e2:
                print(f"Ошибка при извлечении текста с помощью PyPDF2: {str(e2)}")
                raise
        
        # Если текст не извлечен (сканированный PDF), пробуем OCR через Surya
        if not text or len(text.strip()) == 0:
            print("PDF не содержит текста, возможно это сканированный документ. Пробуем OCR через Surya...")
            try:
                # Пробуем использовать pdf2image для конвертации PDF в изображения
                try:
                    from pdf2image import convert_from_bytes
                    from backend.llm_client import recognize_text_from_image_llm_svc
                    from PIL import Image
                    import io
                    
                    # Конвертируем PDF в изображения из bytes
                    images = convert_from_bytes(file_data, dpi=300)
                    print(f"PDF конвертирован в {len(images)} изображений для OCR")
                    
                    # Применяем OCR к каждому изображению через Surya
                    ocr_text = ""
                    ocr_confidence_scores = []
                    for i, image in enumerate(images):
                        print(f"Обрабатываем страницу {i+1}/{len(images)} с помощью Surya OCR...")
                        
                        # Конвертируем PIL Image в bytes
                        img_bytes = io.BytesIO()
                        image.save(img_bytes, format='PNG')
                        img_bytes.seek(0)
                        page_image_data = img_bytes.getvalue()
                        
                        # Вызываем OCR через llm-svc API
                        result = recognize_text_from_image_llm_svc(
                            image_file=page_image_data,
                            filename=f"page_{i+1}.png",
                            languages="ru,en"
                        )
                        
                        if result.get("success", False):
                            page_text = result.get("text", "")
                            page_confidence = result.get("confidence", 50.0)
                            ocr_text += f"\n--- Страница {i+1} ---\n{page_text}\n"
                            ocr_confidence_scores.append(page_confidence)
                        else:
                            print(f"Ошибка OCR для страницы {i+1}: {result.get('error', 'Unknown error')}")
                            ocr_confidence_scores.append(50.0)
                    
                    if ocr_text.strip():
                        text = ocr_text
                        confidence_scores = ocr_confidence_scores
                        print(f"Surya OCR успешно извлек {len(text)} символов из {len(images)} страниц")
                    else:
                        print("Surya OCR не смог извлечь текст из PDF")
                except ImportError:
                    print("Библиотека pdf2image не установлена. Для обработки сканированных PDF установите: pip install pdf2image")
                    print("Также требуется установить poppler: https://github.com/oschwartz10612/poppler-windows/releases")
                except Exception as ocr_error:
                    print(f"Ошибка при OCR обработке PDF через Surya: {ocr_error}")
                    import traceback
                    traceback.print_exc()
            except Exception as e:
                print(f"Не удалось применить OCR к PDF: {e}")
        
        # Вычисляем среднюю уверенность
        avg_confidence = sum(confidence_scores) / len(confidence_scores) if confidence_scores else 0.0
        
        # Создаем слова с уверенностью (для PDF используем среднюю уверенность или 100%)
        confidence_per_word = 100.0 if avg_confidence > 90.0 else avg_confidence
        confidence_info = self._create_confidence_info_for_text(text, confidence_per_word, "pdf")
        confidence_info["pages_processed"] = len(confidence_scores)
        
        return {
            "text": text,
            "confidence_info": confidence_info
        }
    
    def extract_text_from_excel(self, file_path):
        """Извлечение текста из Excel файла"""
        print(f"Извлекаем текст из Excel файла: {file_path}")
        workbook = openpyxl.load_workbook(file_path, data_only=True)
        text_content = []
        
        for sheet_name in workbook.sheetnames:
            sheet = workbook[sheet_name]
            text_content.append(f"Лист: {sheet_name}")
            
            for row in sheet.iter_rows():
                row_values = []
                for cell in row:
                    if cell.value is not None:
                        row_values.append(str(cell.value))
                if row_values:
                    text_content.append("\t".join(row_values))
        
        result = "\n".join(text_content)
        print(f"Извлечено {len(result)} символов из Excel")
        return result
    
    def extract_text_from_txt_bytes(self, file_data: bytes):
        """Извлечение текста из TXT файла из bytes"""
        print(f"Извлекаем текст из TXT файла (размер: {len(file_data)} байт)")
        try:
            # Пробуем декодировать как UTF-8
            result = file_data.decode('utf-8')
            print(f"UTF-8 успешно извлек {len(result)} символов")
            return result
        except UnicodeDecodeError:
            # Если не удалось декодировать как UTF-8, пробуем другие кодировки
            encodings = ['cp1251', 'latin-1', 'koi8-r']
            for encoding in encodings:
                try:
                    result = file_data.decode(encoding)
                    print(f"{encoding} успешно извлек {len(result)} символов")
                    return result
                except UnicodeDecodeError:
                    continue
            
            # Если все кодировки не подошли, возвращаем строковое представление
            result = str(file_data)
            print(f"Бинарный режим извлек {len(result)} символов")
            return result

    def extract_text_from_image_bytes(self, file_data: bytes):
        """Извлечение текста из изображения из bytes с помощью Surya OCR через llm-svc API с информацией об уверенности"""
        print(f"Извлекаем текст из изображения с помощью Surya OCR (размер: {len(file_data)} байт)")
        print(f"DEBUG: Начинаем вызов OCR через llm-svc API...")
        try:
            # Импортируем функцию для работы с OCR через llm-svc
            from backend.llm_client import recognize_text_from_image_llm_svc
            from PIL import Image
            
            # Определяем имя файла на основе формата изображения
            img = Image.open(BytesIO(file_data))
            filename = "image.jpg"
            if img.format:
                filename = f"image.{img.format.lower()}"
            
            print(f"DEBUG: Изображение открыто, формат: {img.format}, размер: {img.size}")
            
            # Вызываем OCR через llm-svc API
            print("Отправляем запрос на распознавание текста через llm-svc...")
            print(f"DEBUG: Вызываем recognize_text_from_image_llm_svc с filename={filename}, languages=ru,en")
            try:
                result = recognize_text_from_image_llm_svc(
                    image_file=file_data,
                    filename=filename,
                    languages="ru,en"
                )
                print(f"DEBUG: OCR вернул результат: success={result.get('success', False)}")
            except Exception as ocr_exception:
                print(f"DEBUG: Исключение при вызове OCR: {ocr_exception}")
                import traceback
                traceback.print_exc()
                raise
            
            # Проверяем результат
            if not result.get("success", False):
                error_msg = result.get("error", "Неизвестная ошибка")
                print(f"Surya OCR вернул ошибку: {error_msg}")
                print(f"ВНИМАНИЕ: OCR не удался, текст не будет извлечен из изображения")
                # Не сохраняем сообщение об ошибке как текст документа
                # Вместо этого возвращаем пустой текст
                return {
                    "text": "",  # Пустой текст вместо сообщения об ошибке
                    "confidence_info": {
                        "confidence": 0.0,
                        "text_length": 0,
                        "file_type": "image",
                        "ocr_available": False,
                        "error": error_msg,
                        "words": []
                    }
                }
            
            # Извлекаем данные из результата
            text = result.get("text", "")
            words_with_confidence = result.get("words", [])
            avg_confidence = result.get("confidence", 0.0)
            word_count = result.get("words_count", 0)
            
            # Если текст не извлечен, возвращаем пустой текст
            if not text.strip():
                print(f"Surya OCR не смог извлечь текст из изображения (текст пустой)")
                # Не сохраняем сообщение об отсутствии текста как текст документа
                return {
                    "text": "",  # Пустой текст
                    "confidence_info": {
                        "confidence": 0.0,
                        "text_length": 0,
                        "file_type": "image",
                        "ocr_available": False,
                        "words": []
                    }
                }
            
            print(f"Surya OCR успешно извлек {len(text)} символов, {word_count} слов, средняя уверенность: {avg_confidence:.2f}%")
            
            return {
                "text": text,
                "confidence_info": {
                    "confidence": avg_confidence,
                    "text_length": len(text),
                    "file_type": "image",
                    "ocr_available": True,
                    "words": words_with_confidence
                }
            }
        except ImportError:
            # Если функция не доступна, возвращаем информацию о файле
            result_text = f"[Изображение. Для распознавания текста требуется доступ к llm-svc API.]"
            print(f"Функция распознавания через llm-svc не доступна, возвращаем описание: {len(result_text)} символов")
            return {
                "text": result_text,
                "confidence_info": {
                    "confidence": 0.0,
                    "text_length": len(result_text),
                    "file_type": "image",
                    "ocr_available": False,
                    "words": []
                }
            }
        except Exception as e:
            error_msg = str(e)
            print(f"Ошибка при обработке изображения через Surya OCR: {error_msg}")
            import traceback
            traceback.print_exc()
            
            # Не сохраняем сообщение об ошибке как текст документа
            # Вместо этого возвращаем пустой текст, но сохраняем информацию об ошибке
            result_text = ""  # Пустой текст вместо сообщения об ошибке
            print(f"ВНИМАНИЕ: OCR не удался, текст не будет извлечен из изображения")
            return {
                "text": result_text,
                "confidence_info": {
                    "confidence": 0.0,
                    "text_length": 0,
                    "file_type": "image",
                    "ocr_available": False,
                    "error": error_msg,
                    "words": []
                }
            }
    
    async def add_document_to_collection_async(self, text, doc_name):
        """Асинхронное добавление документа в коллекцию и обновление векторного хранилища"""
        logger.info(f"Добавление документа '{doc_name}' в коллекцию...")
        logger.info(f"Длина текста: {len(text)} символов")
        
        # Оптимизированное разбиение текста для баланса между контекстом и скоростью
        # Увеличенный размер чанка = меньше чанков = быстрее обработка, но больше контекста в каждом чанке
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1500,  # Увеличили для меньшего количества чанков и большей скорости
            chunk_overlap=200,  # Умеренное перекрытие для сохранения связей
            length_function=len,
        )
        
        chunks = text_splitter.split_text(text)
        print(f"Создано чанков: {len(chunks)}")
        
        # Если нет чанков (пустой текст), создаем хотя бы один минимальный чанк
        if len(chunks) == 0:
            print(f"ВНИМАНИЕ: Нет чанков для документа '{doc_name}', создаем минимальный чанк")
            chunks = [text] if text else [f"[Документ: {doc_name}]"]
        
        # Создаем документы для langchain (для кэша)
        langchain_docs = []
        for i, chunk in enumerate(chunks):
            langchain_docs.append(
                Document(
                    page_content=chunk,
                    metadata={"source": doc_name, "chunk": i}
                )
            )
        
        # Добавляем в общий список документов (кэш)
        self.documents.extend(langchain_docs)
        if doc_name not in self.doc_names:
            self.doc_names.append(doc_name)
        
        # Обновляем кэш структуры документа для быстрого доступа
        self._doc_chunks_cache[doc_name] = [
            {
                "content": doc.page_content,
                "chunk": doc.metadata.get("chunk", i)
            }
            for i, doc in enumerate(langchain_docs)
        ]
        # Сортируем по номеру чанка
        self._doc_chunks_cache[doc_name].sort(key=lambda x: x['chunk'])
        
        print(f"Документ добавлен в кэш. Всего документов: {len(self.documents)}, имен: {len(self.doc_names)}")
        print(f"Кэш структуры обновлен для '{doc_name}': {len(self._doc_chunks_cache[doc_name])} чанков")
        
        # Сохраняем в PostgreSQL + pgvector
        if self.vector_repo and self.document_repo:
            try:
                logger.info(f"Сохранение документа '{doc_name}' в PostgreSQL + pgvector...")
                await self._save_document_to_pgvector(text, doc_name, chunks)
                logger.info(f"Документ '{doc_name}' успешно сохранен в PostgreSQL + pgvector")
                # Устанавливаем флаг vectorstore после успешного сохранения
                self.vectorstore = True
            except Exception as e:
                logger.error(f"ОШИБКА при сохранении документа в PostgreSQL: {str(e)}")
                import traceback
                logger.debug(f"Traceback: {traceback.format_exc()}")
                # Даже при ошибке сохраняем в памяти, но vectorstore остается None
        else:
            logger.warning(f"PostgreSQL недоступен, документ '{doc_name}' сохранен только в памяти")
            # Если pgvector недоступен, но документы есть в памяти, можно использовать их
            if self.doc_names:
                logger.info(f"Документы доступны в памяти: {len(self.doc_names)} документов")
    
    def add_document_to_collection(self, text, doc_name):
        """Добавление документа в коллекцию (синхронная обертка для обратной совместимости)"""
        # Пытаемся получить текущий event loop
        try:
            loop = asyncio.get_event_loop()
            if loop.is_running():
                # Если event loop уже запущен, создаем задачу
                # Но это синхронный метод, поэтому используем run_until_complete с новым loop
                # Или создаем задачу в фоне
                import concurrent.futures
                import threading
                
                # Создаем новый event loop в отдельном потоке для выполнения async операции
                def run_in_thread():
                    new_loop = asyncio.new_event_loop()
                    asyncio.set_event_loop(new_loop)
                    try:
                        new_loop.run_until_complete(self.add_document_to_collection_async(text, doc_name))
                    finally:
                        new_loop.close()
                
                thread = threading.Thread(target=run_in_thread)
                thread.start()
                thread.join()  # Ждем завершения
            else:
                # Если loop не запущен, можем использовать run_until_complete
                loop.run_until_complete(self.add_document_to_collection_async(text, doc_name))
        except RuntimeError:
            # Если нет event loop, создаем новый
            asyncio.run(self.add_document_to_collection_async(text, doc_name))
    
    async def _save_document_to_pgvector(self, text: str, doc_name: str, chunks: List[str]):
        """Сохранение документа в PostgreSQL с векторами"""
        if not self.embeddings:
            logger.error("Модель эмбеддингов не инициализирована")
            return
        
        try:
            # Проверяем, существует ли документ
            document_id = self.filename_to_id.get(doc_name)
            
            # Создаем или обновляем документ в БД
            if document_id is None:
                # Создаем новый документ
                logger.info(f"   Создание нового документа в БД: '{doc_name}'")
                pg_doc = PGDocument(
                    filename=doc_name,
                    content=text,
                    metadata={
                        "confidence_data": self.confidence_data.get(doc_name, {}),
                        "chunks_count": len(chunks)
                    },
                    created_at=datetime.utcnow(),
                    updated_at=datetime.utcnow()
                )
                document_id = await self.document_repo.create_document(pg_doc)
                if document_id:
                    self.filename_to_id[doc_name] = document_id
                    logger.info(f"Документ создан в БД: ID={document_id}")
                else:
                    logger.error("Не удалось создать документ в БД")
                    return
            else:
                # Обновляем существующий документ
                logger.info(f"Обновление существующего документа в БД: ID={document_id}")
                pg_doc = await self.document_repo.get_document(document_id)
                if pg_doc:
                    pg_doc.content = text
                    pg_doc.metadata = {
                        "confidence_data": self.confidence_data.get(doc_name, {}),
                        "chunks_count": len(chunks)
                    }
                    pg_doc.updated_at = datetime.utcnow()
                    # Удаляем старые векторы
                    await self.vector_repo.delete_vectors_by_document(document_id)
            
            # Генерируем эмбеддинги и сохраняем векторы
            logger.info(f"Генерация эмбеддингов для {len(chunks)} чанков...")
            saved_vectors = 0
            for i, chunk in enumerate(chunks):
                try:
                    # Генерируем эмбеддинг для чанка
                    embedding = self.embeddings.embed_query(chunk)
                    
                    # Создаем вектор
                    vector = DocumentVector(
                        document_id=document_id,
                        chunk_index=i,
                        embedding=embedding,
                        content=chunk,
                        metadata={"source": doc_name, "chunk": i}
                    )
                    
                    # Сохраняем вектор в БД
                    vector_id = await self.vector_repo.create_vector(vector)
                    if vector_id:
                        saved_vectors += 1
                        if (i + 1) % 10 == 0 or i == len(chunks) - 1:
                            logger.info(f"Сохранено векторов: {i+1}/{len(chunks)}")
                except Exception as e:
                    logger.warning(f" Ошибка при сохранении вектора {i}: {str(e)}")
                    continue
            
            if saved_vectors == len(chunks):
                logger.info(f"Все {saved_vectors} векторов успешно сохранены в pgvector")
            else:
                logger.warning(f"Сохранено {saved_vectors}/{len(chunks)} векторов")
            
        except Exception as e:
            logger.error(f"ОШИБКА при сохранении документа в pgvector: {str(e)}")
            import traceback
            logger.debug(f"Traceback: {traceback.format_exc()}")
            raise
    
    def update_vectorstore(self):
        """Обновление или создание векторного хранилища (для обратной совместимости)"""
        # При использовании pgvector векторы сохраняются сразу при добавлении документа
        # Этот метод оставлен для обратной совместимости
        if self.vector_repo:
            print("Используется pgvector - векторы сохраняются автоматически при добавлении документов")
            self.vectorstore = True
        else:
            print("ВНИМАНИЕ: pgvector недоступен, векторное хранилище не обновлено")
            self.vectorstore = None
    
    async def query_documents_async(self, query, k=2):
        """Асинхронный поиск релевантных документов по запросу"""
        logger.info(f"Поиск релевантных документов для запроса: '{query[:50]}...'")
        
        # Проверяем наличие vector_repo (pgvector) вместо vectorstore
        if not self.vector_repo:
            logger.warning("pgvector недоступен (vector_repo is None)")
            return "Векторное хранилище не инициализировано или pgvector недоступен"
        
        # Проверяем, есть ли загруженные документы
        if not self.doc_names or len(self.doc_names) == 0:
            logger.warning("Нет загруженных документов для поиска")
            return "Нет загруженных документов"
        
        if not self.embeddings:
            logger.error("Модель эмбеддингов не инициализирована")
            return "Модель эмбеддингов не инициализирована"
        
        try:
            logger.info(f"Начинаем поиск документов для запроса: '{query[:100]}...'")
            logger.info(f"Параметры поиска: k={k}, doc_names={self.doc_names}")
            # Используем pgvector для поиска
            results = await self._query_documents_async(query, k)
            if isinstance(results, list):
                logger.info(f"Найдено {len(results)} релевантных документов через pgvector")
                if len(results) == 0:
                    logger.warning("Векторный поиск не вернул результатов. Возможно, документы не сохранены в pgvector.")
            elif isinstance(results, str):
                logger.error(f"Ошибка поиска (строка): {results}")
            return results
        except Exception as e:
            logger.error(f"Ошибка при поиске по документам: {str(e)}")
            import traceback
            logger.error(f"Traceback: {traceback.format_exc()}")
            return f"Ошибка при поиске по документам: {str(e)}"
    
    def query_documents(self, query, k=2):
        """Поиск релевантных документов по запросу (синхронная обертка для обратной совместимости)"""
        try:
            loop = asyncio.get_running_loop()
            # Если loop запущен, это ошибка - нужно использовать async версию
            logger.error("query_documents вызван из async контекста. Используйте query_documents_async()")
            return "Ошибка: используйте async версию метода"
        except RuntimeError:
            # Нет запущенного loop, можем использовать asyncio.run
            return asyncio.run(self.query_documents_async(query, k))
    
    async def _query_documents_async(self, query: str, k: int = 2):
        """Асинхронный поиск релевантных документов через pgvector"""
        logger.debug(f"   Выполняем векторный поиск через pgvector с k={k}...")
        
        # Генерируем эмбеддинг для запроса
        query_embedding = self.embeddings.embed_query(query)
        logger.debug(f"   Эмбеддинг запроса сгенерирован (размерность: {len(query_embedding)})")
        
        # Выполняем поиск в pgvector
        results = await self.vector_repo.similarity_search(query_embedding, limit=k)
        logger.debug(f"   pgvector вернул {len(results)} результатов")
        
        # Преобразуем результаты в нужный формат
        formatted_results = []
        for vector, similarity in results:
            # Получаем имя файла по document_id
            doc_name = None
            for filename, doc_id in self.filename_to_id.items():
                if doc_id == vector.document_id:
                    doc_name = filename
                    break
            
            if doc_name is None:
                # Пытаемся получить из БД
                pg_doc = await self.document_repo.get_document(vector.document_id)
                if pg_doc:
                    doc_name = pg_doc.filename
                    self.filename_to_id[doc_name] = vector.document_id
            
            result = {
                "content": vector.content,
                "source": doc_name or f"document_{vector.document_id}",
                "chunk": vector.chunk_index,
                "similarity": similarity
            }
            formatted_results.append(result)
            logger.debug(f"{result['source']}, чанк {result['chunk']}, similarity: {similarity:.4f}")
        
        return formatted_results
    
    def get_document_list(self):
        """Получение списка загруженных документов"""
        print(f"get_document_list вызван. Документы: {self.doc_names}")
        return self.doc_names
    
    def get_image_paths(self):
        """
        Получение списка путей к изображениям для мультимодальной модели
        Возвращает список путей к локальным файлам (временные файлы, скачанные из MinIO при необходимости)
        """
        image_paths_list = []
        print(f"DEBUG get_image_paths: image_paths = {self.image_paths}")
        
        for filename, path_info in self.image_paths.items():
            print(f"DEBUG get_image_paths: обрабатываем {filename}, path_info = {path_info}")
            
            if isinstance(path_info, dict):
                # Если есть file_data, создаем временный файл
                if "file_data" in path_info:
                    try:
                        import tempfile
                        # Создаем временный файл
                        suffix = os.path.splitext(filename)[1] or ".jpg"
                        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=suffix)
                        temp_file.write(path_info["file_data"])
                        temp_file.close()
                        temp_path = temp_file.name
                        print(f"DEBUG get_image_paths: создан временный файл для {filename}: {temp_path}")
                        image_paths_list.append(temp_path)
                    except Exception as e:
                        print(f"ERROR get_image_paths: не удалось создать временный файл для {filename}: {e}")
                        image_paths_list.append(None)
                # Если есть путь, используем его
                elif "path" in path_info:
                    image_paths_list.append(path_info.get("path"))
                else:
                    print(f"WARNING get_image_paths: нет file_data или path для {filename}")
                    image_paths_list.append(None)
            else:
                # Обратная совместимость: просто путь
                image_paths_list.append(path_info)
        
        print(f"get_image_paths вызван. Изображения: {image_paths_list}")
        return image_paths_list
    
    def get_image_minio_info(self, filename):
        """
        Получение информации о MinIO объекте для изображения
        
        Returns:
            dict: {"minio_object": str, "minio_bucket": str} или None
        """
        if filename in self.image_paths:
            path_info = self.image_paths[filename]
            if isinstance(path_info, dict):
                return {
                    "minio_object": path_info.get("minio_object"),
                    "minio_bucket": path_info.get("minio_bucket")
                }
        return None
    
    def clear_documents(self):
        """Очистка коллекции документов"""
        print("Очищаем коллекцию документов...")
        self.documents = []
        self.doc_names = []
        self.vectorstore = None
        self.confidence_data = {}
        self.image_paths = {}
        self._doc_chunks_cache = {}  # Очищаем кэш
        self.filename_to_id = {}  # Очищаем маппинг
        
        # Очищаем из PostgreSQL (опционально, можно оставить данные в БД)
        # Если нужно очистить БД, раскомментируйте:
        # if self.document_repo:
        #     # Удаляем все документы из БД
        #     asyncio.run(self._clear_all_documents_from_db())
        
        print("Коллекция документов очищена")
        return "Коллекция документов очищена"
    
    def get_confidence_report_data(self):
        """Получение данных для отчета об уверенности с процентами над словами"""
        if not self.confidence_data:
            return {
                "total_documents": 0,
                "documents": [],
                "average_confidence": 0.0,
                "formatted_texts": []
            }
        
        documents = []
        formatted_texts = []
        total_confidence = 0.0
        total_weighted_confidence = 0.0
        total_words = 0
        
        for filename, info in self.confidence_data.items():
            words = info.get("words", [])
            
            # Форматируем текст с процентами над словами
            formatted_lines = []
            current_line = []
            
            for word_info in words:
                word = word_info.get("word", "")
                conf = word_info.get("confidence", 0.0)
                
                # Пропускаем пустые слова
                if not word:
                    continue
                
                # Форматируем каждое слово с процентом над ним
                formatted_word = f"{conf:.0f}%\n{word}"
                current_line.append(formatted_word)
                
                # Добавляем перенос строки после каждого слова для читаемости
                if len(current_line) >= 10:  # Примерно 10 слов на строку
                    formatted_lines.append("  ".join(current_line))
                    current_line = []
            
            # Добавляем оставшиеся слова
            if current_line:
                formatted_lines.append("  ".join(current_line))
            
            formatted_text = "\n".join(formatted_lines)
            
            # Вычисляем среднюю уверенность для документа
            doc_avg_confidence = info.get("confidence", 0.0)
            if words:
                doc_avg_confidence = sum(w.get("confidence", 0.0) for w in words) / len(words)
            
            documents.append({
                "filename": filename,
                "confidence": doc_avg_confidence,
                "text_length": info.get("text_length", 0),
                "file_type": info.get("file_type", "unknown"),
                "words_count": len(words)
            })
            
            formatted_texts.append({
                "filename": filename,
                "formatted_text": formatted_text,
                "words": words
            })
            
            total_confidence += doc_avg_confidence
            if words:
                total_weighted_confidence += sum(w.get("confidence", 0.0) for w in words)
                total_words += len(words)
        
        # Вычисляем общую среднюю уверенность
        avg_confidence = total_confidence / len(documents) if documents else 0.0
        
        # Вычисляем итоговую уверенность по всем словам
        overall_confidence = total_weighted_confidence / total_words if total_words > 0 else avg_confidence
        
        return {
            "total_documents": len(documents),
            "documents": documents,
            "average_confidence": avg_confidence,
            "overall_confidence": overall_confidence,
            "total_words": total_words,
            "formatted_texts": formatted_texts
        }
    
    def process_query(self, query, agent_function):
        """Обработка запроса с контекстом документов для LLM"""
        print(f"Обрабатываем запрос: {query}")
        print(f"Векторное хранилище: {self.vectorstore is not None}")
        print(f"Количество документов: {len(self.documents)}")
        print(f"Имена документов: {self.doc_names}")
        
        if not self.vectorstore:
            return "Нет загруженных документов. Пожалуйста, загрузите документы перед выполнением запроса."
        
        try:
            # Получаем релевантные документы
            docs = self.query_documents(query)
            print(f"Найдено релевантных фрагментов: {len(docs) if isinstance(docs, list) else 'ошибка'}")
            
            if isinstance(docs, str):  # Если возникла ошибка
                print(f"Ошибка при поиске документов: {docs}")
                return docs
            
            # Формируем контекст из найденных документов
            context = "Контекст из документов:\n\n"
            for i, doc in enumerate(docs):
                context += f"Фрагмент {i+1} (из документа '{doc['source']}'):\n{doc['content']}\n\n"
            
            print(f"Контекст сформирован, длина: {len(context)} символов")
            
            # Подготавливаем запрос для LLM с инструкциями и контекстом
            prompt = f"""На основе предоставленного контекста ответь на вопрос пользователя. 
Если информации в контексте недостаточно, укажи это.
Отвечай только на основе информации из контекста. Не придумывай информацию.

{context}

Вопрос пользователя: {query}

Ответ:"""
            
            print("Отправляем запрос к LLM...")
            # Отправляем запрос к LLM
            response = agent_function(prompt)
            print(f"Получен ответ от LLM, длина: {len(response)} символов")
            return response
            
        except Exception as e:
            print(f"Ошибка при обработке запроса: {str(e)}")
            return f"Ошибка при обработке запроса: {str(e)}"
    
    def remove_document(self, filename):
        """Удалить конкретный документ по имени файла"""
        print(f"Удаляем документ: {filename}")
        print(f"До удаления - self.doc_names: {self.doc_names}")
        print(f"До удаления - self.documents: {len(self.documents)}")
        print(f"До удаления - self.vectorstore доступен: {self.vectorstore is not None}")
        
        try:
            # Находим индекс документа
            if filename not in self.doc_names:
                print(f"Документ {filename} не найден")
                return False
            
            # Удаляем документ из списка имен
            index = self.doc_names.index(filename)
            self.doc_names.pop(index)
            print(f"Документ {filename} удален из списка имен")
            
            # Удаляем информацию об уверенности
            if filename in self.confidence_data:
                del self.confidence_data[filename]
                print(f"Информация об уверенности для {filename} удалена")
            
            # Удаляем путь к изображению, если это изображение
            if filename in self.image_paths:
                del self.image_paths[filename]
                print(f"Путь к изображению для {filename} удален")
            
            # Удаляем из кэша структуры документа
            if filename in self._doc_chunks_cache:
                del self._doc_chunks_cache[filename]
                print(f"Кэш структуры для {filename} удален")
            
            # Удаляем ВСЕ чанки этого документа из списка документов
            # Ищем все документы с этим именем и удаляем их
            documents_to_remove = []
            for i, doc in enumerate(self.documents):
                if doc.metadata.get("source") == filename:
                    documents_to_remove.append(i)
            
            # Удаляем чанки в обратном порядке, чтобы индексы не сдвигались
            for i in reversed(documents_to_remove):
                self.documents.pop(i)
            
            print(f"Удалено чанков документа {filename}: {len(documents_to_remove)}")
            print(f"После удаления - self.doc_names: {self.doc_names}")
            print(f"После удаления - self.documents: {len(self.documents)}")
            
            # Удаляем из PostgreSQL + pgvector
            if self.document_repo and filename in self.filename_to_id:
                try:
                    document_id = self.filename_to_id[filename]
                    # Удаляем векторы и документ асинхронно
                    try:
                        loop = asyncio.get_running_loop()
                        # Если loop запущен, используем create_task или run в отдельном потоке
                        import threading
                        
                        def run_async():
                            new_loop = asyncio.new_event_loop()
                            asyncio.set_event_loop(new_loop)
                            try:
                                new_loop.run_until_complete(self.vector_repo.delete_vectors_by_document(document_id))
                                new_loop.run_until_complete(self.document_repo.delete_document(document_id))
                            finally:
                                new_loop.close()
                        
                        thread = threading.Thread(target=run_async)
                        thread.start()
                        thread.join()
                    except RuntimeError:
                        # Нет запущенного loop
                        asyncio.run(self.vector_repo.delete_vectors_by_document(document_id))
                        asyncio.run(self.document_repo.delete_document(document_id))
                    
                    # Удаляем из маппинга
                    del self.filename_to_id[filename]
                    print(f"Документ {filename} удален из PostgreSQL + pgvector")
                except Exception as e:
                    print(f"Ошибка при удалении документа из PostgreSQL: {str(e)}")
                    import traceback
                    traceback.print_exc()
            
            # Обновляем флаг vectorstore
            if not self.doc_names:
                print("Нет документов, очищаем vectorstore")
                self.vectorstore = None
            else:
                self.vectorstore = True if self.vector_repo else None
            
            print(f"Документ {filename} успешно удален. Осталось документов: {len(self.doc_names)}")
            return True
            
        except Exception as e:
            print(f"Ошибка при удалении документа {filename}: {str(e)}")
            import traceback
            traceback.print_exc()
            return False 
    
    async def get_document_context_async(self, query, k=2, include_all_chunks=None, max_context_length=30000):
        """
        Получение контекста документов для запроса с оптимизацией для скорости
        
        Args:
            query: Запрос пользователя
            k: Количество релевантных фрагментов для векторного поиска
            include_all_chunks: Если None - автоматически определяет стратегию по типу запроса
                              Если True - включает все чанки документа
                              If False - только релевантные фрагменты
            max_context_length: Максимальная длина контекста в символах (по умолчанию 30000)
        """
        print(f"Получаем контекст документов для запроса: '{query}'")
        logger.info(f"get_document_context вызван для запроса: '{query[:100]}...'")
        logger.info(f"vector_repo доступен: {self.vector_repo is not None}")
        logger.info(f"doc_names: {self.doc_names}")
        logger.info(f"Количество документов: {len(self.doc_names) if self.doc_names else 0}")
        
        # Проверяем наличие vector_repo и документов
        if not self.vector_repo:
            logger.warning("pgvector недоступен (vector_repo is None)")
            print("pgvector недоступен (vector_repo is None)")
            return None
        
        if not self.doc_names or len(self.doc_names) == 0:
            logger.warning("Нет загруженных документов")
            print("Нет загруженных документов")
            return None
        
        try:
            import time
            start_time = time.time()
            
            # Автоматическое определение стратегии по типу запроса
            if include_all_chunks is None:
                query_lower = query.lower()
                # Запросы, требующие полного контекста
                full_context_keywords = [
                    'саммари', 'summary', 'краткое содержание', 'обзор', 'резюме',
                    'по всему документу', 'весь документ', 'всего документа',
                    'перескажи', 'опиши весь', 'расскажи о документе', 'структура документа'
                ]
                # Конкретные вопросы - используем только релевантные фрагменты
                is_full_context_request = any(keyword in query_lower for keyword in full_context_keywords)
                include_all_chunks = is_full_context_request
                print(f"Автоматически определен режим: {'ПОЛНЫЙ КОНТЕКСТ' if include_all_chunks else 'РЕЛЕВАНТНЫЕ ФРАГМЕНТЫ'}")
            else:
                print(f"Режим: {'ПОЛНЫЙ КОНТЕКСТ' if include_all_chunks else 'РЕЛЕВАНТНЫЕ ФРАГМЕНТЫ'}")
            
            if include_all_chunks:
                # ОПТИМИЗИРОВАННЫЙ РЕЖИМ: Используем векторный поиск только для определения релевантных документов,
                # затем быстро получаем ВСЕ чанки этих документов из кэша или БД
                
                # Быстрый векторный поиск для определения, какие документы релевантны
                docs = await self.query_documents_async(query, k=min(k, 5))  # Нужно только для определения документов
                print(f"Векторный поиск найденных документов: {len(docs) if isinstance(docs, list) else 'ошибка'}")
                
                if isinstance(docs, str):  # Если возникла ошибка
                    print(f"Ошибка при поиске документов: {docs}")
                    return None
                
                # Собираем уникальные имена документов
                doc_names_found = set()
                for doc in docs:
                    if isinstance(doc, dict) and 'source' in doc:
                        doc_names_found.add(doc['source'])
                
                # Если документы не найдены через поиск, используем все загруженные документы
                if not doc_names_found and self.doc_names:
                    doc_names_found = set(self.doc_names)
                    print(f"Документы не найдены через поиск, используем все загруженные: {list(doc_names_found)}")
                
                # ОПТИМИЗИРОВАННАЯ СТРАТЕГИЯ: Для полного контекста используем выборку ключевых чанков
                # вместо всех чанков для ускорения обработки
                all_chunks = []
                for doc_name in doc_names_found:
                    if doc_name in self._doc_chunks_cache:
                        cached_chunks = self._doc_chunks_cache[doc_name]
                        total_chunks = len(cached_chunks)
                        
                        # Выбираем ключевые чанки: начало, конец, и равномерно распределенные
                        selected_chunks = []
                        
                        # Всегда включаем первый чанк
                        if cached_chunks:
                            selected_chunks.append(cached_chunks[0])
                        
                        # Всегда включаем последний чанк
                        if len(cached_chunks) > 1:
                            selected_chunks.append(cached_chunks[-1])
                        
                        # Равномерно распределяем остальные чанки (примерно 15-20 для баланса)
                        target_chunks = min(18, total_chunks)  # Оптимальное количество для скорости
                        if total_chunks > 2:
                            step = max(1, total_chunks // target_chunks)
                            for i in range(step, total_chunks - 1, step):
                                if cached_chunks[i] not in selected_chunks:
                                    selected_chunks.append(cached_chunks[i])
                        
                        # Сортируем по номеру чанка
                        selected_chunks.sort(key=lambda x: x['chunk'])
                        
                        # Преобразуем в нужный формат
                        for chunk_data in selected_chunks:
                            all_chunks.append({
                                "content": chunk_data["content"],
                                "source": doc_name,
                                "chunk": chunk_data["chunk"]
                            })
                        
                        print(f"Выбрано {len(selected_chunks)} ключевых чанков из {total_chunks} для документа '{doc_name}'")
                    else:
                        # Fallback: если кэш не обновлен
                        print(f"Кэш не найден для '{doc_name}', используем fallback")
                        doc_chunks = []
                        for doc_item in self.documents:
                            if doc_item.metadata.get("source") == doc_name:
                                doc_chunks.append({
                                    "content": doc_item.page_content,
                                    "source": doc_name,
                                    "chunk": doc_item.metadata.get("chunk", 0)
                                })
                        doc_chunks.sort(key=lambda x: x['chunk'])
                        all_chunks.extend(doc_chunks)
                        self._doc_chunks_cache[doc_name] = [
                            {"content": c["content"], "chunk": c["chunk"]} 
                            for c in doc_chunks
                        ]
                
                # Сортируем все чанки по документу и номеру чанка
                all_chunks.sort(key=lambda x: (x['source'], x['chunk']))
                
                # Формируем контекст с ограничением длины
                context_parts = []
                current_length = 0
                chunks_added = 0
                
                for chunk in all_chunks:
                    if chunk['chunk'] == 0:
                        fragment = f"[НАЧАЛО ДОКУМЕНТА '{chunk['source']}']\n{chunk['content']}"
                    else:
                        fragment = f"[Чанк {chunk['chunk']} из '{chunk['source']}']\n{chunk['content']}"
                    
                    # Проверяем ограничение длины
                    if current_length + len(fragment) > max_context_length:
                        print(f"Достигнуто ограничение длины ({max_context_length} символов). Добавлено {chunks_added} из {len(all_chunks)} чанков.")
                        break
                    
                    context_parts.append(fragment)
                    current_length += len(fragment)
                    chunks_added += 1
                
                context = "\n\n".join(context_parts) + "\n\n"
                
                elapsed_time = time.time() - start_time
                print(f"Контекст сформирован за {elapsed_time:.2f}с: {len(context)} символов, {chunks_added}/{len(all_chunks)} чанков")
                
            else:
                # ОПТИМИЗИРОВАННЫЙ РЕЖИМ: Для конкретных вопросов используем релевантные фрагменты + начало документа
                # Увеличиваем k для получения большего количества релевантных фрагментов
                k = max(k, 8)  # Минимум 8 релевантных фрагментов
                docs = await self.query_documents_async(query, k=k)
                print(f"Найдено релевантных фрагментов: {len(docs) if isinstance(docs, list) else 'ошибка'}")
                
                if isinstance(docs, str):
                    print(f"Ошибка при поиске документов: {docs}")
                    return None
                
                # Собираем уникальные документы
                doc_names_found = set()
                for doc in docs:
                    if isinstance(doc, dict) and 'source' in doc:
                        doc_names_found.add(doc['source'])
                
                # Добавляем первый чанк каждого документа для контекста
                context_parts = []
                added_chunks = set()
                
                # Сначала добавляем первые чанки документов
                for doc_name in doc_names_found:
                    if doc_name in self._doc_chunks_cache and self._doc_chunks_cache[doc_name]:
                        first_chunk = self._doc_chunks_cache[doc_name][0]
                        chunk_key = (doc_name, first_chunk['chunk'])
                        if chunk_key not in added_chunks:
                            context_parts.append(f"[НАЧАЛО ДОКУМЕНТА '{doc_name}']\n{first_chunk['content']}")
                            added_chunks.add(chunk_key)
                
                # Затем добавляем релевантные фрагменты
                for doc in docs:
                    chunk_key = (doc['source'], doc['chunk'])
                    if chunk_key not in added_chunks:
                        context_parts.append(f"Фрагмент (из документа '{doc['source']}', чанк {doc['chunk']}):\n{doc['content']}")
                        added_chunks.add(chunk_key)
                
                # Ограничиваем длину контекста
                context = "\n\n".join(context_parts)
                if len(context) > max_context_length:
                    # Обрезаем до максимальной длины, сохраняя начало
                    context = context[:max_context_length]
                    context += "\n\n[Контекст обрезан для оптимизации скорости]"
                
                context += "\n\n"
                
                elapsed_time = time.time() - start_time
                print(f"Контекст сформирован за {elapsed_time:.2f}с: {len(context)} символов, {len(context_parts)} фрагментов")
            
            return context
            
        except Exception as e:
            print(f"Ошибка при получении контекста документов: {str(e)}")
            import traceback
            traceback.print_exc()
            return None
    
    def get_document_context(self, query, k=2, include_all_chunks=None, max_context_length=30000):
        """Синхронная обертка для get_document_context_async (для обратной совместимости)"""
        try:
            loop = asyncio.get_running_loop()
            # Если loop запущен, это ошибка - нужно использовать async версию
            logger.error("get_document_context вызван из async контекста. Используйте get_document_context_async()")
            return None
        except RuntimeError:
            # Нет запущенного loop, можем использовать asyncio.run
            return asyncio.run(self.get_document_context_async(query, k, include_all_chunks, max_context_length))
            
        except Exception as e:
            print(f"Ошибка при получении контекста документов: {str(e)}")
            import traceback
            traceback.print_exc()
            return None
    
