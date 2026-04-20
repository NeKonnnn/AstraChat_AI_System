import logging
import re
from io import BytesIO
from typing import Any, Dict, List, Optional

import httpx

from app.core.config import get_settings

logger = logging.getLogger(__name__)

try:
    import docx

    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import PyPDF2

    PYPDF2_AVAILABLE = True
except ImportError:
    PYPDF2_AVAILABLE = False

try:
    import fitz  # PyMuPDF

    PYMUPDF_AVAILABLE = True
except ImportError:
    PYMUPDF_AVAILABLE = False

try:
    import pdfplumber

    PDFPLUMBER_AVAILABLE = True
except ImportError:
    PDFPLUMBER_AVAILABLE = False

try:
    import openpyxl

    OPENPYXL_AVAILABLE = True
except ImportError:
    OPENPYXL_AVAILABLE = False

try:
    import xlrd

    XLRD_AVAILABLE = True
except ImportError:
    XLRD_AVAILABLE = False

try:
    from pdf2image import convert_from_bytes as pdf_convert_from_bytes

    PDF2IMAGE_AVAILABLE = True
except ImportError:
    PDF2IMAGE_AVAILABLE = False

try:
    from PIL import Image

    PIL_AVAILABLE = True
except ImportError:
    PIL_AVAILABLE = False


def _create_confidence_info_for_text(text: str, confidence_per_word: float, file_type: str) -> Dict[str, Any]:
    """Создаёт структуру confidence_info, совместимую с backend."""
    tokens = (text or "").split()
    words_with_confidence = [{"word": token, "confidence": float(confidence_per_word)} for token in tokens]
    avg_confidence = confidence_per_word
    return {
        "confidence": avg_confidence,
        "text_length": len(text or ""),
        "file_type": file_type,
        "words": words_with_confidence,
    }


async def _call_ocr_service(image_bytes: bytes, filename: str, languages: str = "ru,en") -> Dict[str, Any]:
    """Вызов OCR (Surya) по URL из config.yml """
    settings = get_settings()
    ocr_url = settings.ocr.url.rstrip("/")
    timeout = settings.ocr.timeout

    mime = "image/jpeg"
    lower = filename.lower()
    if lower.endswith(".png"):
        mime = "image/png"

    files = {"file": (filename, BytesIO(image_bytes), mime)}
    data = {"languages": languages}

    try:
        req_timeout = httpx.Timeout(timeout, connect=10.0, read=timeout, write=10.0)
        async with httpx.AsyncClient(timeout=req_timeout) as client:
            resp = await client.post(f"{ocr_url}/v1/ocr", files=files, data=data, headers={"Accept": "application/json"})
            resp.raise_for_status()
            return resp.json()
    except Exception as e:
        logger.error("Ошибка OCR при обращении к %s: %s", ocr_url, e)
        raise


def extract_text_from_docx(file_data: bytes) -> str:
    if not DOCX_AVAILABLE:
        raise RuntimeError("python-docx не установлен")
    doc = docx.Document(BytesIO(file_data))
    parts = []
    for para in doc.paragraphs:
        parts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


def _pdf_magic_bytes(data: bytes) -> bool:
    return bool(data and len(data) >= 5 and data[:5] == b"%PDF-")


def _normalize_extracted_text(text: str) -> str:
    """Нормализация извлечённого текста перед чанкингом/эмбеддингом."""
    t = (text or "").replace("\r\n", "\n").replace("\r", "\n")
    # Мягкие переносы и разрывы слов на стыках строк из PDF
    t = t.replace("\u00ad", "")
    t = re.sub(r"(?<=\w)-\n(?=\w)", "", t)
    # Убираем «рваные» пробелы, но сохраняем абзацы
    t = re.sub(r"[ \t]+", " ", t)
    t = re.sub(r"\n{3,}", "\n\n", t)
    return t.strip()


def _looks_like_low_quality_pdf_text(text: str, n_pages: int) -> bool:
    """Эвристика плохого текстового слоя PDF (часто сканы/битые шрифты)."""
    t = (text or "").strip()
    if not t:
        return True
    pages = max(int(n_pages or 1), 1)
    min_chars_expected = max(120, pages * 80)
    if len(t) < min_chars_expected:
        return True
    letters = sum(1 for ch in t if ch.isalpha())
    digits = sum(1 for ch in t if ch.isdigit())
    spaces = sum(1 for ch in t if ch.isspace())
    total = max(len(t), 1)
    alpha_ratio = letters / total
    # Очень мало букв при длинном тексте → вероятный "мусорный" слой
    if alpha_ratio < 0.35 and (digits + spaces) / total > 0.45:
        return True
    # Слишком много одиночных токенов/обрывков
    toks = re.findall(r"\w+", t, re.UNICODE)
    if toks:
        short_ratio = sum(1 for w in toks if len(w) <= 2) / len(toks)
        if short_ratio > 0.72:
            return True
    return False


async def extract_text_from_pdf_bytes(file_data: bytes) -> Dict[str, Any]:
    """
    PyMuPDF + pdfplumber + PyPDF2 — выбирается самый длинный извлечённый текст;
    при пустом или слишком коротком слое относительно числа страниц — OCR (pdf2image + ocr-service).
    """
    logger.info("PDF: извлечение текста, размер=%s байт", len(file_data))
    candidates: List[tuple] = []
    n_pages = 0

    if PYMUPDF_AVAILABLE:
        try:
            doc = fitz.open(stream=file_data, filetype="pdf")
            if getattr(doc, "needs_pass", False):
                try:
                    doc.authenticate("")
                except Exception:
                    pass
            n_pages = doc.page_count
            parts: List[str] = []
            for i in range(n_pages):
                parts.append(doc.load_page(i).get_text() or "")
            doc.close()
            t = _normalize_extracted_text("\n".join(parts))
            if t.strip():
                candidates.append(("pymupdf", t))
                logger.info("PDF: PyMuPDF извлёк %s символов, страниц=%s", len(t), n_pages)
        except Exception as e:
            logger.warning("PDF: PyMuPDF ошибка: %s", e)

    text_pb_parts: List[str] = []
    if PDFPLUMBER_AVAILABLE:
        try:
            with pdfplumber.open(BytesIO(file_data)) as pdf:
                n_pages = max(n_pages, len(pdf.pages))
                for page in pdf.pages:
                    page_text = page.extract_text() or ""
                    text_pb_parts.append(page_text)
            text_pb = _normalize_extracted_text("\n".join(text_pb_parts))
            if text_pb.strip():
                candidates.append(("pdfplumber", text_pb))
                logger.info("PDF: pdfplumber извлёк %s символов", len(text_pb))
        except Exception as e:
            logger.warning("PDF: pdfplumber ошибка: %s", e)

    if PYPDF2_AVAILABLE:
        try:
            reader = PyPDF2.PdfReader(BytesIO(file_data))
            n_pages = max(n_pages, len(reader.pages))
            text_p2_parts: List[str] = []
            for page in reader.pages:
                text_p2_parts.append(page.extract_text() or "")
            text_p2 = _normalize_extracted_text("\n".join(text_p2_parts))
            if text_p2.strip():
                candidates.append(("pypdf2", text_p2))
                logger.info("PDF: PyPDF2 извлёк %s символов", len(text_p2))
        except Exception as e2:
            logger.warning("PDF: PyPDF2 ошибка: %s", e2)

    text = ""
    confidence_scores: List[float] = []
    if candidates:
        label, text = max(candidates, key=lambda x: len(x[1].strip()))
        logger.info("PDF: для индекса выбран движок «%s» (%s символов)", label, len(text.strip()))

    weak_layer = _looks_like_low_quality_pdf_text(text, n_pages)

    need_ocr = (not text.strip() or weak_layer) and PDF2IMAGE_AVAILABLE and PIL_AVAILABLE
    if need_ocr:
        if weak_layer:
            logger.info(
                "PDF: текстовый слой низкого качества (%s симв., страниц=%s) — пробуем OCR",
                len(text.strip()),
                max(n_pages, 1),
            )
        else:
            logger.info("PDF: текстовый слой пуст — rasterize+OCR (poppler + ocr-service)")
        try:
            try:
                images = pdf_convert_from_bytes(file_data, dpi=300)
            except Exception as e_dpi:
                logger.warning("PDF: pdf2image dpi=300 не удалось (%s), пробуем dpi=150", e_dpi)
                images = pdf_convert_from_bytes(file_data, dpi=150)
            logger.info("PDF: растеризация в %s страниц для OCR", len(images))

            ocr_text = ""
            ocr_confidence_scores: List[float] = []
            for i, image in enumerate(images):
                buf = BytesIO()
                image.save(buf, format="PNG")
                buf.seek(0)
                page_image_data = buf.getvalue()
                result = await _call_ocr_service(page_image_data, f"page_{i+1}.png", languages="ru,en")
                if result.get("success"):
                    page_text = result.get("text", "")
                    ocr_text += f"\n--- Страница {i+1} ---\n{page_text}\n"
                    ocr_confidence_scores.append(float(result.get("confidence", 50.0) or 50.0))
                else:
                    logger.warning("PDF OCR: страница %s: %s", i + 1, result.get("error", "Unknown"))
                    ocr_confidence_scores.append(50.0)

            ocr_text = _normalize_extracted_text(ocr_text)
            if ocr_text.strip():
                if len(ocr_text.strip()) >= len(text.strip()):
                    text = ocr_text
                    confidence_scores = ocr_confidence_scores
                    logger.info("PDF: использован OCR-текст (%s символов)", len(text.strip()))
                else:
                    logger.info(
                        "PDF: OCR не дал больше текста (%s vs %s симв.) — оставляем текстовый слой",
                        len(ocr_text.strip()),
                        len(text.strip()),
                    )
            else:
                logger.warning("PDF: OCR не вернул текст")
        except Exception as e:
            logger.warning("PDF: OCR недоступен или сбой (poppler/pdf2image/ocr-service): %s", e)

    if text.strip() and not confidence_scores:
        confidence_scores = [95.0] * max(1, n_pages or 1)

    avg_confidence = sum(confidence_scores) / len(confidence_scores) if confidence_scores else 0.0
    confidence_per_word = 100.0 if avg_confidence > 90.0 else avg_confidence
    text = _normalize_extracted_text(text)
    confidence_info = _create_confidence_info_for_text(text or "", confidence_per_word, "pdf")
    confidence_info["pages_processed"] = len(confidence_scores) if confidence_scores else (n_pages or 0)

    if not (text or "").strip():
        logger.error(
            "PDF: итоговый текст пуст (страниц≈%s). Нужны рабочие PyMuPDF/pdfplumber, для сканов — poppler в образе и "
            "доступный ocr-service (SVC_OCR_URL).",
            n_pages,
        )

    return {
        "text": text or "",
        "confidence_info": confidence_info,
        "file_type": "pdf",
        "pages": int(confidence_info.get("pages_processed") or 0),
    }


def extract_text_from_xls_xlrd(file_data: bytes) -> str:
    """Старый формат Excel .xls (не .xlsx)."""
    if not XLRD_AVAILABLE:
        raise RuntimeError("xlrd не установлен")
    book = xlrd.open_workbook(file_contents=file_data)
    parts: List[str] = []
    for si in range(book.nsheets):
        sh = book.sheet_by_index(si)
        parts.append(f"Лист: {sh.name}")
        for ri in range(sh.nrows):
            row = sh.row(ri)
            vals = [str(c.value) for c in row if c.value not in ("", None)]
            if vals:
                parts.append("\t".join(vals))
    return "\n".join(parts)


def extract_text_from_xlsx(file_data: bytes) -> str:
    if not OPENPYXL_AVAILABLE:
        raise RuntimeError("openpyxl не установлен")
    workbook = openpyxl.load_workbook(BytesIO(file_data), data_only=True)
    parts = []
    for sheet_name in workbook.sheetnames:
        sheet = workbook[sheet_name]
        parts.append(f"Лист: {sheet_name}")
        for row in sheet.iter_rows():
            row_vals = [str(c.value) for c in row if c.value is not None]
            if row_vals:
                parts.append("\t".join(row_vals))
    return "\n".join(parts)


def extract_text_from_txt(file_data: bytes) -> str:
    for enc in ("utf-8", "cp1251", "latin-1", "koi8-r"):
        try:
            return file_data.decode(enc)
        except UnicodeDecodeError:
            continue
    return file_data.decode("utf-8", errors="replace")


async def extract_text_from_image_bytes(file_data: bytes) -> Dict[str, Any]:
    """Извлечение текста из изображения (OCR). Вызов идёт в ocr-service (Surya)."""
    print(f"Извлекаем текст из изображения с помощью Surya OCR (размер: {len(file_data)} байт)")
    if not PIL_AVAILABLE:
        result_text = "[Изображение. Для распознавания текста требуется Pillow и доступ к ocr-service.]"
        return {
            "text": result_text,
            "confidence_info": _create_confidence_info_for_text(result_text, 0.0, "image"),
        }

    img = Image.open(BytesIO(file_data)).convert("RGB")
    filename = "image.jpg"
    if img.format:
        filename = f"image.{img.format.lower()}"

    print(f"DEBUG: Изображение открыто, формат: {img.format}, размер: {img.size}")

    # Увеличиваем маленькие изображения, как в backend
    min_side = 1024
    w, h = img.size
    if max(w, h) < min_side and max(w, h) > 0:
        scale = min_side / max(w, h)
        new_w = max(1, int(round(w * scale)))
        new_h = max(1, int(round(h * scale)))
        img = img.resize((new_w, new_h), Image.Resampling.LANCZOS)
        print(f"DEBUG: Изображение увеличено для OCR: {w}x{h} -> {new_w}x{new_h}")

    buf = BytesIO()
    img.save(buf, format="PNG")
    image_to_send = buf.getvalue()

    try:
        result = await _call_ocr_service(image_to_send, filename, languages="ru,en")
    except Exception as e:
        print(f"Ошибка OCR: {e}")
        return {
            "text": "",
            "confidence_info": {
                "confidence": 0.0,
                "text_length": 0,
                "file_type": "image",
                "ocr_available": False,
                "error": str(e),
                "words": [],
            },
        }

    if not result.get("success"):
        error_msg = result.get("error", "Неизвестная ошибка")
        print(f"Surya OCR вернул ошибку: {error_msg}")
        return {
            "text": "",
            "confidence_info": {
                "confidence": 0.0,
                "text_length": 0,
                "file_type": "image",
                "ocr_available": False,
                "error": error_msg,
                "words": [],
            },
        }

    text = result.get("text", "") or ""
    words = result.get("words", []) or []
    avg_confidence = float(result.get("confidence", 0.0) or 0.0)

    if not text.strip():
        print("Surya OCR не смог извлечь текст из изображения (текст пустой)")
        return {
            "text": "",
            "confidence_info": {
                "confidence": 0.0,
                "text_length": 0,
                "file_type": "image",
                "ocr_available": False,
                "words": [],
            },
        }

    print(f"Surya OCR успешно извлек {len(text)} символов, {len(words)} слов, средняя уверенность: {avg_confidence:.2f}%")

    confidence_info = {
        "confidence": avg_confidence,
        "text_length": len(text),
        "file_type": "image",
        "ocr_available": True,
        "words": words,
    }
    return {"text": text, "confidence_info": confidence_info}


async def parse_document(file_data: bytes, filename: str) -> Optional[Dict[str, Any]]:
    """
    По расширению файла выбираем парсер и возвращаем структуру:
    {"text": str, "confidence_info": dict}
    """
    name = (filename or "").lower()
    # Корректно извлекаем только сам суффикс (".docx", ".pdf" и т.п.)
    dot = name.rfind(".")
    ext = name[dot:] if dot != -1 else ""
    if ext != ".pdf" and _pdf_magic_bytes(file_data):
        logger.info(
            "parse_document: файл «%s» без .pdf, но с сигнатурой PDF — парсим как PDF",
            filename or "?",
        )
        return await extract_text_from_pdf_bytes(file_data)
    if ext == ".docx":
        if not DOCX_AVAILABLE:
            return None
        text = extract_text_from_docx(file_data)
        return {
            "text": text,
            "confidence_info": _create_confidence_info_for_text(text, 100.0, "docx"),
        }
    if ext == ".pdf":
        return await extract_text_from_pdf_bytes(file_data)
    if ext == ".xlsx":
        if not OPENPYXL_AVAILABLE:
            logger.warning("Парсинг .xlsx: openpyxl не установлен")
            return None
        try:
            text = extract_text_from_xlsx(file_data)
        except Exception as e:
            logger.error("Ошибка парсинга .xlsx: %s", e)
            return None
        return {
            "text": text,
            "confidence_info": _create_confidence_info_for_text(text, 100.0, "excel"),
        }
    if ext == ".xls":
        if not XLRD_AVAILABLE:
            logger.warning("Парсинг .xls: xlrd не установлен")
            return None
        try:
            text = extract_text_from_xls_xlrd(file_data)
        except Exception as e:
            logger.error("Ошибка парсинга .xls: %s", e)
            return None
        return {
            "text": text,
            "confidence_info": _create_confidence_info_for_text(text, 100.0, "excel"),
        }
    if ext == ".txt":
        text = extract_text_from_txt(file_data)
        return {
            "text": text,
            "confidence_info": _create_confidence_info_for_text(text, 100.0, "txt"),
        }
    if ext in (".jpg", ".jpeg", ".png", ".webp"):
        return await extract_text_from_image_bytes(file_data)
    return None
